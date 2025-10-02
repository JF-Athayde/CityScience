from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import os
import datetime

OUTPUT_DIR = 'output_images'
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ----- Exemplo de dados (substitua com seus dados reais) -----
np.random.seed(42)
months = pd.date_range(start='2024-01-01', periods=12, freq='M')
temp = 30 + 20 * np.sin(np.linspace(0, 2 * np.pi, 12)) + np.random.normal(0, 3, 12)
hum = 60 + 20 * np.cos(np.linspace(0, 2 * np.pi, 12)) + np.random.normal(0, 5, 12)

# --- ADDED: four additional synthetic series useful for the bulletin ---
aqi = 50 + 15 * np.sin(np.linspace(0.5, 2.5 * np.pi, 12)) + np.random.normal(0, 6, 12)
pm25 = 8 + 12 * np.abs(np.sin(np.linspace(0, 2 * np.pi, 12))) + np.random.normal(0, 2, 12)
dry_flag = (np.sin(np.linspace(-0.3, 2.0, 12)) > 0).astype(float)
fire_hotspots = np.random.poisson(lam=1.5 + 4.0 * dry_flag)  # integer counts, higher in dry months
precip = 120 * np.clip(np.cos(np.linspace(0, 2 * np.pi, 12)), 0, 1) + np.random.normal(0, 10, 12)

df = pd.DataFrame({
    'date': months,
    'temperature_C': temp,
    'humidity_%': hum,
    'AQI': aqi,
    'pm25_ugm3': pm25,
    'fire_hotspots': fire_hotspots,
    'precip_mm': precip
})

# ----- Funções utilitárias -----

def save_plot(df, column, ylabel, title, path):
    plt.figure(figsize=(8, 4.5))
    plt.plot(df['date'], df[column], marker='o')
    plt.title(title)
    plt.xlabel('Date')
    plt.ylabel(ylabel)
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(path)
    plt.close()

# ----- Estatísticas -----
# UPDATED: return month names for max/min
def get_stats(series, dates):
    idx_max = int(series.idxmax())
    idx_min = int(series.idxmin())
    return {
        'max': series.max(),
        'min': series.min(),
        'mean': series.mean(),
        'std': series.std(),
        'amplitude': series.max() - series.min(),
        'max_month': dates.iloc[idx_max].strftime('%B'),
        'min_month': dates.iloc[idx_min].strftime('%B')
    }

# use updated get_stats for all series
temp_stats = get_stats(df['temperature_C'], df['date'])
hum_stats = get_stats(df['humidity_%'], df['date'])
aqi_stats = get_stats(df['AQI'], df['date'])
pm25_stats = get_stats(df['pm25_ugm3'], df['date'])
fires_stats = get_stats(df['fire_hotspots'], df['date'])
precip_stats = get_stats(df['precip_mm'], df['date'])

# ----- Salva gráficos -----

temp_img = os.path.join(OUTPUT_DIR, 'temperature.png')
hum_img = os.path.join(OUTPUT_DIR, 'humidity.png')

# --- ADDED image paths for new plots ---
aqi_img = os.path.join(OUTPUT_DIR, 'aqi.png')
pm25_img = os.path.join(OUTPUT_DIR, 'pm25.png')
fires_img = os.path.join(OUTPUT_DIR, 'fire_hotspots.png')
precip_img = os.path.join(OUTPUT_DIR, 'precipitation.png')

save_plot(df, 'temperature_C', 'Temperature (°C)', 'Temperature - Series 1', temp_img)
save_plot(df, 'humidity_%', 'Humidity (%)', 'Humidity - Series 1', hum_img)

# --- Save the four new plots ---
save_plot(df, 'AQI', 'AQI', 'Air Quality Index (Series 1)', aqi_img)
save_plot(df, 'pm25_ugm3', 'PM2.5 (µg/m³)', 'PM2.5 (Series 1)', pm25_img)
# fire hotspots (counts) - use same helper
save_plot(df, 'fire_hotspots', 'Hotspots (count)', 'Fire Hotspots (Series 1)', fires_img)
save_plot(df, 'precip_mm', 'Precipitation (mm)', 'Precipitation (Series 1)', precip_img)

# ----- Monta o documento -----

doc = Document()

# Fonte padrão
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

# Título
h = doc.add_heading('', level=1)
h_run = h.add_run('Environmental Bulletin 2025')
h_run.bold = True
h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Subtítulo / metadados
meta = doc.add_paragraph()
meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
meta.add_run('Automatic Zoom    City Science\n').italic = True
meta.add_run('Data source: City Science sensors and external datasets (e.g. MODIS). Location:\n')
meta.add_run('Fortaleza - CE, Brazil. Extracted from: Nasa Database').italic = True

doc.add_paragraph()  # espaçamento

# Resumo de temperatura (use computed months)
p_temp = doc.add_paragraph()
p_temp.add_run('Temperature shows a clear seasonal pattern. ').bold = True
p_temp.add_run(
    f"Series 1 peaked at {temp_stats['max']:.1f} °C in {temp_stats['max_month']} and dropped to a low of {temp_stats['min']:.1f} °C in {temp_stats['min_month']}, an amplitude of {temp_stats['amplitude']:.1f} °C. "
)
p_temp.add_run(
    f"The average was {temp_stats['mean']:.1f} °C (σ={temp_stats['std']:.1f}), indicating notable differences between seasons."
)

# Insere gráfico de temperatura
doc.add_paragraph()
doc.add_picture(temp_img, width=Inches(6))

# Resumo de umidade (use computed months)
p_hum = doc.add_paragraph()
p_hum.add_run('Humidity measurements follow recurring cycles with distinct highs and lows. ').bold = True
p_hum.add_run(
    f"Series 1 reached a maximum of {hum_stats['max']:.1f} % in {hum_stats['max_month']} and a minimum of {hum_stats['min']:.1f} % in {hum_stats['min_month']}. "
)
p_hum.add_run(
    f"With a mean of {hum_stats['mean']:.1f} % (σ={hum_stats['std']:.1f}), these values help identify periods relevant for agriculture and water management."
)

# Insere gráfico de umidade
doc.add_paragraph()
doc.add_picture(hum_img, width=Inches(6))

# --- ADDED: Air Quality section ---
doc.add_paragraph()
p_aqi = doc.add_paragraph()
p_aqi.add_run('Air Quality (AQI) summary. ').bold = True
p_aqi.add_run(
    f"Series 1 peaked at {aqi_stats['max']:.1f} AQI in {aqi_stats['max_month']} and fell to {aqi_stats['min']:.1f} AQI in {aqi_stats['min_month']}. "
)
p_aqi.add_run(
    f"The mean was {aqi_stats['mean']:.1f} AQI (σ={aqi_stats['std']:.1f}), indicating periods that should be correlated with local emission sources and meteorology."
)
doc.add_paragraph()
doc.add_picture(aqi_img, width=Inches(6))

# --- ADDED: PM2.5 section ---
doc.add_paragraph()
p_pm = doc.add_paragraph()
p_pm.add_run('PM2.5 concentration summary. ').bold = True
p_pm.add_run(
    f"Series 1 reached {pm25_stats['max']:.1f} µg/m³ in {pm25_stats['max_month']} and a minimum of {pm25_stats['min']:.1f} µg/m³ in {pm25_stats['min_month']}. "
)
p_pm.add_run(
    f"Mean {pm25_stats['mean']:.1f} µg/m³ (σ={pm25_stats['std']:.1f}). Elevated PM2.5 episodes may impact public health and merit targeted mitigation."
)
doc.add_paragraph()
doc.add_picture(pm25_img, width=Inches(6))

# --- ADDED: Fire hotspots section ---
doc.add_paragraph()
p_fire = doc.add_paragraph()
p_fire.add_run('Fire hotspots summary. ').bold = True
p_fire.add_run(
    f"Series 1 shows a peak of {int(fires_stats['max'])} hotspots in {fires_stats['max_month']} and a low of {int(fires_stats['min'])} in {fires_stats['min_month']}. "
)
p_fire.add_run(
    f"Average {fires_stats['mean']:.1f} hotspots (σ={fires_stats['std']:.1f}). These patterns help identify high‑risk months for fire prevention."
)
doc.add_paragraph()
doc.add_picture(fires_img, width=Inches(6))

# --- ADDED: Precipitation section ---
doc.add_paragraph()
p_prec = doc.add_paragraph()
p_prec.add_run('Precipitation summary. ').bold = True
p_prec.add_run(
    f"Series 1 recorded a maximum monthly total of {precip_stats['max']:.1f} mm in {precip_stats['max_month']} and a minimum of {precip_stats['min']:.1f} mm in {precip_stats['min_month']}. "
)
p_prec.add_run(
    f"Mean {precip_stats['mean']:.1f} mm (σ={precip_stats['std']:.1f}). Precipitation seasonality is important for hydrology and fire-risk analysis."
)
doc.add_paragraph()
doc.add_picture(precip_img, width=Inches(6))

# Rodapé de data

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
footer_p.add_run(f'Generated: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

# Salva o documento
output_docx = 'Boletim_Ambiental_2025.docx'
doc.save(output_docx)

print(f'Document generated: {output_docx}')
print(f'Images saved in: {OUTPUT_DIR}')
